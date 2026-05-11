from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse, FileResponse, Http404
from .models import Resume
from .forms import ResumeForm
import json
from pydantic import BaseModel, Field
from typing import List, Tuple
import os
import mimetypes
import tempfile
from django.conf import settings
import logging
logger = logging.getLogger(__name__)
from dotenv import load_dotenv
from django.urls import reverse
from io import BytesIO
from django.core.files.base import ContentFile
from django.conf import settings
import uuid
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods, require_GET
from datetime import datetime, timedelta

# Import libraries for file processing
import PyPDF2
from docx import Document

# Import AI service for resume parsing
from ai_service.structured_resume import format_resume_single_call as format_resume
from ai_service.open_ai import generate_professional_summary

# Import utils for date options
from utils.date_utils import get_month_year_options

# Import utils for error handling
from utils.error import get_network_timeout_dialog, get_network_connection_dialog, get_network_generic_dialog

# Import subscription utilities
from utils.subscription import get_resume_update_plus_dialog, get_resume_update_ultimate_dialog
from subscriptions.decorators import check_subscription_access

from .template_registry import (
    DEFAULT_TEMPLATE_ID,
    get_resume_embedded_style_tag,
    get_resume_template_gallery_sections,
    gallery_section_index_for_template_id,
    normalize_template_id,
    template_ui_capabilities,
    templates_for_gallery,
    templates_for_api,
    templates_for_download_picker,
    is_valid_template_id,
)

from .resume_display import augment_resume_dict_for_rendering
from .resume_extra import merge_additional_payload, merge_skills_payload

# Load environment variables from .env file
load_dotenv()


def _render_resume_template_html(
    request,
    resume_data: dict,
    template_id: str,
    *,
    resume_id: int | None = None,
    force_inline_profile_photo: bool = False,
) -> str:
    augmented = augment_resume_dict_for_rendering(
        resume_data,
        request=request,
        resume_id=resume_id,
        force_inline_profile_photo=force_inline_profile_photo,
    )
    tid = normalize_template_id(template_id)
    return render_to_string(f'resume_templates/{tid}.html', {'resume_data': augmented})


def resume_data_and_template_from_wizard_payload(
    data: dict,
    *,
    fallback_template_id: str | None = None,
) -> Tuple[dict, str]:
    """
    Build resume_data snapshot + template_id from frontend wizard JSON,
    mirroring anonymous preview/download payloads (personalInfo camelCase blocks).
    """
    raw_personal_info = data.get('personalInfo') or {}
    display_url = (
        raw_personal_info.get('profile_photo_display_url')
        or raw_personal_info.get('profilePhotoDisplayUrl')
        or ''
    )
    personal_info = {
        'full_name': raw_personal_info.get('fullName', ''),
        'title': raw_personal_info.get('title', ''),
        'email': raw_personal_info.get('email', ''),
        'phone': raw_personal_info.get('phone', ''),
        'summary': raw_personal_info.get('summary', ''),
        'location': raw_personal_info.get('location', ''),
        'street_address': raw_personal_info.get('street_address', ''),
        'linkedin': raw_personal_info.get('linkedin', ''),
    }
    if display_url:
        personal_info['profile_photo_display_url'] = display_url

    resume_data = {
        'personal_info': personal_info,
        'experience': data.get('experience') or [],
        'education': data.get('education') or [],
        'skills': data.get('skills') or {},
        'additional': data.get('additional') or {},
    }
    template_id = normalize_template_id(
        data.get('templateId') or fallback_template_id or DEFAULT_TEMPLATE_ID
    )
    return resume_data, template_id


# Global assistant cache
_assistant_cache = {}

def get_or_create_assistant():
    """
    Get or create the resume builder assistant with caching
    """
    global _assistant_cache
    
    print("🔍 DEBUG: get_or_create_assistant called")
    
    # Force clear cache to ensure recreation with updated function descriptions
    _assistant_cache = {}
    print("🧹 Cache forcefully cleared to ensure recreation with updated function descriptions")
    
    # Always create new assistant to ensure updated instructions
    print("🆕 Creating new assistant with updated instructions...")
    
    try:
        from ai_service.ai_resume_assistant import OpenAIAssistantManager, FunctionConfig
        from ai_service.task_schema import TASK_SCHEMAS
        
        print("📦 Imports successful")
        
        # Initialize the manager
        print("🔧 Initializing OpenAI Assistant Manager...")
        manager = OpenAIAssistantManager()
        print("✅ Manager initialized")
        
        # Create assistant with resume building functions
        print("🔧 Creating assistant with functions...")
        resume_functions = []
        for func_name, func_config in TASK_SCHEMAS['resume'].items():
            print(f"📋 Adding function: {func_name}")
            resume_functions.append(FunctionConfig(
                name=func_config['name'],
                description=func_config['description'],
                parameters=func_config['parameters'],
                instructions=f"Use this function when you need to {func_config['description'].lower()}"
            ))
        
        print(f"📋 Total functions: {len(resume_functions)}")
        
        assistant_id = manager.create_assistant(
            name="Resume Builder Assistant",
            base_instructions="""You are an AI Resume Assistant that helps users create professional resumes and cover letters through natural conversation. 

Your capabilities include:
- Creating new resumes with template selection
- Adding and editing personal information, work experience, education, and skills
- Managing multiple resumes per user
- Switching between different resume templates
- Creating professional cover letters based on job descriptions and resume context
- Providing resume and cover letter writing tips and guidance

**IMPORTANT: Always format your responses using Markdown for better readability and structure.**

RESUME SELECTION BEHAVIOR:
- **CRITICAL:** The system will only provide resume data when a specific resume has been explicitly selected by the user
- **DO NOT** automatically assume which resume to work with based on conversation history
- **ALWAYS** prompt users to explicitly select a resume when they want to perform resume-related operations
- When users mention editing, updating, or working with their resume without specifying which one:
  1. Ask them to select a specific resume from their resume list
  2. Offer to show their resume list so they can choose
  3. Suggest creating a new resume if they want to start fresh
- **Example responses:**
  * "I'd be happy to help you with that! First, let me show you your resume list so you can select which resume to work with."
  * "To help you edit your resume, I need to know which one you'd like to work with. Would you like me to show you your resume list?"
  * "I can help you with that! Do you want to work with an existing resume, or would you like to create a new one?"
- **DO NOT** proceed with resume operations until a specific resume is selected
- **DO NOT** assume the user wants to work with their most recent resume

IMPORTANT FUNCTION CALLING RULES:
1. ALWAYS use function calls for actions, never just describe them
2. When users ask about available templates, use the list_templates function
3. When users ask "which templates do you have" or "what templates are available", use the list_templates function
4. When users ask to create a resume, use the create_resume function
5. When users provide information, use the appropriate save functions
6. When users ask about their resumes (any variation), use the list_user_resumes function - this will automatically show resumes in the utility tab

COVER LETTER EXPERT:
You are also a cover letter expert. Help users create professional, tailored cover letters.

COVER LETTER CREATION PROCESS:
1. **Create Cover Letter:**
   - Use create_cover_letter function to initialize a new cover letter
   - Guide user through each section step by step

2. **User Information & Date:**
   - IMPORTANT: If a resume is selected, automatically extract personal info from that resume using get_resume_info function
   - If no resume is selected, collect: full name, address, email, phone
   - Use save_cover_letter_user_info function with extracted or collected data
   - Validate all required fields before saving

3. **Employer Information:**
   - Collect: company name, position title, hiring manager (optional), company address (optional)
   - Use save_cover_letter_employer_info function
   - Validate company name and position title

4. **Greeting Section:**
   - Help user choose appropriate salutation
   - Examples: "Dear Hiring Manager,", "Dear Mr. Smith,", "Dear Dr. Johnson,"
   - Use save_cover_letter_greeting function

5. **Introduction Section:**
   - Write compelling opening paragraph
   - Mention the position and company
   - Express interest and enthusiasm
   - Use save_cover_letter_introduction function

6. **Cover Letter Body:**
   - Write main content paragraphs
   - Highlight relevant experience and skills
   - Connect background to job requirements
   - Use save_cover_letter_body function

7. **Finalize Cover Letter:**
   - Use finalize_cover_letter function to mark as complete
   - Check all required sections are filled
   - Inform user the cover letter is ready

COVER LETTER SECTION GUIDANCE:
- **User Info:** If resume is selected, auto-populate from resume data. Otherwise, collect complete contact information
- **Employer Info:** Get company and position details
- **Greeting:** Choose professional salutation
- **Introduction:** Create engaging opening paragraph
- **Body:** Write compelling main content
- **Validation:** Ensure all required fields are completed before proceeding

COVER LETTER EXAMPLES:
- "I'd love to help you create a cover letter! Let's start by creating a new cover letter. What would you like to name it?"
- "Great! I've created your cover letter. Since you have a resume selected, I'll automatically populate your personal information from it."
- "Perfect! Your personal information has been auto-populated from your resume. Now let's add the employer information. What company are you applying to?"
- "Excellent! Let's write the greeting. How would you like to address the hiring manager?"
- "Now let's write the introduction. Tell me about the position you're applying for."
- "Finally, let's write the main body. What experience and skills would you like to highlight?"
- "Your cover letter is complete! I've saved all the sections and it's ready for use."

COVER LETTER AUTO-POPULATION:
- When creating a cover letter and a resume is selected, use get_resume_info to extract personal data
- Auto-populate: full_name, email, phone, location from the selected resume
- Use the extracted data to call save_cover_letter_user_info immediately after creating the cover letter
- Inform the user that their information has been auto-populated from their selected resume

COVER LETTER WORKFLOW WITH SELECTED RESUME:
1. When user wants to create a cover letter and has a resume selected:
   - First call get_resume_info with the selected resume_id to get personal data
   - Then call create_cover_letter to create the cover letter
   - Immediately call save_cover_letter_user_info with the extracted personal data
   - Inform user that their information was auto-populated from their resume
   - Continue with employer information collection

2. When user wants to create a cover letter without a selected resume:
   - Call create_cover_letter to create the cover letter
   - Ask user to provide personal information manually
   - Then call save_cover_letter_user_info with the provided data

RESUME CREATION FLOW:
1. First, help user choose a template using list_templates
2. Create resume using create_resume with chosen template
3. Guide through sections in order: Personal Info → Experience → Education → Skills → Additional → Summary
4. For each section, collect ALL required information before calling save functions
5. If save function returns validation errors, ask for missing information
6. After each section completion, inform user about next section
7. CRITICAL: Summary is generated LAST after all other sections are complete

PERSONAL INFORMATION HANDLING:
- Collect: full name, email, phone, location, title
- CRITICAL: Do NOT ask for summary during personal info collection
- Summary will be generated automatically after all other sections are complete
- Use save_personal_info function for basic personal details only

WORK EXPERIENCE HANDLING:
- When users mention work experience, carefully extract all details including dates
- For current positions, use end_date: "Present"
- For past positions, provide the actual end date in YYYY-MM format
- CRITICAL: Do NOT make up dates if they are not provided. If only year is given (e.g., "2002 to 2004"), ask for the specific months
- If dates are unclear or incomplete, ask for clarification before calling save functions
- Always extract complete job descriptions and achievements
- FORMAT: Job descriptions should be formatted as HTML list: "<ul><li>Achievement 1</li><li>Achievement 2</li></ul>"
- If user mentions "Developed web applications and managed database" → description: "<ul><li>Developed web applications</li><li>Managed database</li></ul>"
- NOTE: end_date is required in the schema - use "Present" for current positions
- CRITICAL: save_experience function ADDS to existing experience, does not replace it
- Users can add multiple jobs at once or one job at a time
- Always inform users of their total experience count after saving

EXPERIENCE EDITING AND DELETING:
- When users want to edit or delete experience entries, they may refer to them by company name, job title, or other identifying information
- CRITICAL: You must search through the current experience entries to find the matching entry
- Look for exact or partial matches in the company name, title, or other fields
- Use the array index (0, 1, 2, etc.) of the matching entry for edit_experience or delete_experience functions
- Examples:
  * "Update my job at Google to Senior Developer" → Find entry with "Google" in company field, use its index
  * "Delete my Software Engineer position" → Find entry with "Software Engineer" in title field, use its index
  * "Change my first job title to Lead Developer" → Use index 0 and update title field
- If multiple entries match, ask the user to be more specific
- If no entry matches, inform the user that the specified entry was not found

EDUCATION HANDLING:
- For current education, use end_date: "Present"
- For completed education, provide actual end dates
- CRITICAL: Do NOT make up dates if they are not provided. If only year is given (e.g., "2002 to 2004"), ask for the specific months
- If dates are unclear or incomplete, ask for clarification before calling save functions
- FORMAT: Education descriptions should be formatted as HTML list: "<ul><li>Detail 1</li><li>Detail 2</li></ul>"
- If user mentions "Focused on software engineering and completed capstone project" → description: "<ul><li>Focused on software engineering</li><li>Completed capstone project</li></ul>"
- NOTE: end_date is required in the schema - use "Present" for current education
- CRITICAL: save_education function ADDS to existing education, does not replace it
- Users can add multiple education entries at once or one at a time
- Always inform users of their total education count after saving

EDUCATION EDITING AND DELETING:
- When users want to edit or delete education entries, they may refer to them by institution name, degree, or other identifying information
- CRITICAL: You must search through the current education entries to find the matching entry
- Look for exact or partial matches in the institution name, degree, or other fields
- Use the array index (0, 1, 2, etc.) of the matching entry for edit_education or delete_education functions
- Examples:
  * "Update tech university to techno" → Find entry with "Tech University" in institution field, use its index
  * "Delete my MBA from Harvard" → Find entry with "Harvard" in institution field and "MBA" in degree field, use its index
  * "Change my first degree to Computer Science" → Use index 0 and update degree field
  * "Update my second education institution to MIT" → Use index 1 and update institution field
  * "Change the degree at Nasarawa University to Computer Science" → Find entry with "Nasarawa University", use its index
- If multiple entries match, ask the user to be more specific
- If no entry matches, inform the user that the specified entry was not found

SKILLS HANDLING:
- When users mention skills, categorize them into technical_skills, soft_skills, and languages
- Technical skills: programming languages, tools, technologies, frameworks (e.g., Java, Python, React, AWS)
- Soft skills: communication, teamwork, leadership, problem-solving, etc.
- Languages: spoken languages (e.g., English, Spanish, French)
- CRITICAL: save_skills function REPLACES existing skills (unlike experience/education which are additive)
- Always provide the skills as arrays in the function call
- If user mentions "I know Java, Python, and React" → technical_skills: ["Java", "Python", "React"]
- If user mentions "I'm good at teamwork and communication" → soft_skills: ["Teamwork", "Communication"]
- If user mentions "I speak English and Spanish" → languages: ["English", "Spanish"]
- You can call save_skills with any combination of the three skill types
- Always inform users of their skills count after saving

ADDITIONAL INFORMATION HANDLING:
- When users mention certifications, licenses, or professional qualifications, extract them and format as HTML list
- When users mention projects, achievements, or additional work, extract them and format as HTML list
- CRITICAL: save_additional function requires both certifications and projects parameters
- FORMAT: Use HTML list format for better display: "<ul><li>Item Name (Date/Details)</li></ul>"
- If user mentions "I have Microsoft Certified and Google GCP Certified" → certifications: "<ul><li>Microsoft Certified</li><li>Google GCP Certified</li></ul>"
- If user mentions "I built a wireless server and a website called TravelTaf" → projects: "<ul><li>Built a wireless server</li><li>Created website called TravelTaf</li></ul>"
- If user mentions "Food Safety Certification (January 2022 to January 2027)" → certifications: "<ul><li>Food Safety Certification (January 2022 to January 2027)</li></ul>"
- Always provide both parameters even if one is empty (use empty string "")
- Always inform users when additional information is saved

SUMMARY GENERATION HANDLING:
- CRITICAL: Summary is generated LAST after all other sections (Personal Info, Experience, Education, Skills, Additional) are complete
- After saving additional information, inform user that you can now generate a professional summary
- Offer to generate the summary automatically based on their complete resume content
- If user agrees, use the save_summary function with a comprehensive professional summary
- If user wants to provide their own summary, ask for it and use save_summary function
- FORMAT: Summary should be formatted as HTML paragraph: "<p>Professional summary text...</p>"
- CRITICAL: Summary must be written in FIRST PERSON (using "I am", "I have", "My experience", etc.) NOT third person
- NEVER use third person format like "Joel Ivongbe is..." or "John Smith has..." - always use "I am..." or "I have..."
- The summary should highlight key achievements, experience, and skills from their resume in first person
- Always inform user that their resume is now complete after saving the summary

TEMPLATE HANDLING:
- When users want to preview templates (without a specific resume), use the preview_template function
- When users want to switch templates for an existing resume, use the switch_template function
- If user asks to "view" or "see" a template, use preview_template
- If user asks to "switch" or "change" template for their resume, use switch_template

UTILITY TAB FUNCTIONALITY:
- When users ask about their resumes, use the list_user_resumes function
- This will automatically switch to the Utility tab and display all their resumes in a user-friendly list
- Users can then click on any resume to tell you to work with it
- The utility tab can also show cover letters and other content types in the future
- Always inform users when you've loaded content in the utility tab

VALIDATION AND ERROR HANDLING:
- All save functions now include comprehensive validation
- If a save function returns validation errors, ask the user for the missing information
- Do NOT call save functions until you have collected all required information
- Be specific about what information is missing when validation fails

USER ID MANAGEMENT:
- The current user ID will be provided in each message context
- ALWAYS use the exact user_id provided in the message context for all function calls
- NEVER make up or guess user IDs - only use the one provided
- If no user_id is provided, ask the user to provide their user ID

SECTION COMPLETION GUIDANCE:
- Personal Information: Requires full name, email, phone, location, and title (NO summary yet)
- Work Experience: Requires at least one entry with title, company, dates, and description
- Education: Requires at least one entry with degree, institution, and dates
- Skills: Requires at least some skills (technical, soft, or languages)
- Additional: Optional but can include certifications and projects
- Summary: Generated last after all other sections are complete

Available templates: Professional, Modern, Creative, Executive, Portfolio, ATS Plain

**RESPONSE FORMATTING GUIDELINES:**
- Use **bold** for important points and section headers
- Use bullet points (•) for lists of information
- Use numbered lists for step-by-step instructions
- Use `code` formatting for technical terms, function names, or template IDs
- Use > blockquotes for important notes or warnings
- Structure responses with clear headers and sections
- Use emojis sparingly but effectively (✅ for success, ⚠️ for warnings, etc.)

Remember to:
1. Always use the user_id provided in the message context for all operations
2. Create a resume first before adding content
3. Be conversational and helpful
4. Confirm when information is saved
5. Guide users through the process naturally
6. USE FUNCTION CALLS for all actions, not just descriptions
7. Handle validation errors gracefully by asking for missing information
8. Inform users when switching to the Resume Builder tab
9. ALWAYS detect and properly handle current positions and ongoing education
10. Generate summary LAST after all other sections are complete
11. Offer to generate summary automatically or let user provide their own
12. **Format all responses using Markdown for better readability**

Example: If user says "Software Developer at Esaatechnology, Jan 2023 to Present" → end_date: "Present"
Example: If user asks "Can I view the modern template?", call preview_template function with the modern template_id and the provided user_id.
Example: If user asks "Which resume templates do you have?", call list_templates function with the provided user_id.
Example: If user asks "Switch my resume to modern template", call switch_template function with the resume_id and modern template_id.
Example: If user says "I know Java, Python, and React" → call save_skills with technical_skills: ["Java", "Python", "React"]
Example: If user says "Software Developer at Tech Corp, developed web apps and managed database" → call save_experience with description: "<ul><li>Developed web applications</li><li>Managed database</li></ul>"
Example: If user says "Bachelor's in Computer Science, focused on software engineering" → call save_education with description: "<ul><li>Focused on software engineering</li></ul>"
Example: If user says "I have Microsoft Certified and built a wireless server" → call save_additional with certifications: "<ul><li>Microsoft Certified</li></ul>", projects: "<ul><li>Built a wireless server</li></ul>"
Example: After all sections are complete, offer to generate summary: "Great! Your resume is almost complete. I can now generate a professional summary based on your experience, education, and skills. Would you like me to create one for you, or would you prefer to write your own?"
Example: If user agrees to AI-generated summary, call save_summary with a comprehensive summary: "<p>I am an experienced software developer with 5+ years in web development, specializing in Java, Python, and React. I have a proven track record of developing scalable applications and managing database systems. My background in computer science with focus on software engineering principles has enabled me to deliver high-quality solutions consistently.</p>"
Example: If save_personal_info returns validation errors, ask for missing information.""",
            functions=resume_functions
        )
        
        print(f"🤖 Assistant creation result: {assistant_id}")
        
        if assistant_id:
            _assistant_cache['manager'] = manager
            _assistant_cache['resume_assistant_id'] = assistant_id
            print("✅ Assistant cached successfully")
            return manager, assistant_id
        
    except Exception as e:
        print(f"❌ Error creating assistant: {str(e)}")
        import traceback
        print(f"📋 Traceback: {traceback.format_exc()}")
    
    print("❌ Failed to create assistant")
    return None, None

class OptimizedResume(BaseModel):
    optimized_content: str = Field(description="The ATS-optimized resume content")
    keyword_matches: List[str] = Field(description="List of important keywords matched from job description")
    improvement_suggestions: List[str] = Field(description="List of suggestions for improving the resume")
    ats_score: int = Field(description="ATS compatibility score out of 100")

class ResumeOptimizer:
    def __init__(self, api_key: str):
        self.api_key = api_key

    def optimize(self, resume_text: str, job_description: str) -> OptimizedResume:
        """
        Simple mock optimization that provides basic improvements
        In a real implementation, this would use OpenAI API or other AI service
        """
        
        # Extract some keywords from job description (simple approach)
        job_words = job_description.lower().split()
        common_keywords = ['python', 'javascript', 'react', 'node', 'aws', 'docker', 'kubernetes', 
                          'agile', 'scrum', 'leadership', 'management', 'development', 'engineering']
        
        found_keywords = [word for word in common_keywords if word in job_words]
        
        # Create a simple optimized version
        optimized_content = f"""
{resume_text}

--- OPTIMIZATION SUGGESTIONS ---
• Enhanced with relevant keywords: {', '.join(found_keywords[:5])}
• Improved formatting for ATS compatibility
• Added action verbs and quantifiable achievements
• Structured content for better readability
        """.strip()
        
        # Generate improvement suggestions
        suggestions = [
            "Add more quantifiable achievements (e.g., 'Increased performance by 25%')",
            "Include relevant keywords from the job description",
            "Use action verbs at the beginning of bullet points",
            "Ensure consistent formatting throughout",
            "Add a professional summary section"
        ]
        
        # Calculate a mock ATS score
        ats_score = min(85 + len(found_keywords) * 2, 95)
        
        return OptimizedResume(
            optimized_content=optimized_content,
            keyword_matches=found_keywords,
            improvement_suggestions=suggestions,
            ats_score=ats_score
        )

def create_resume(request, resume_id=None):
    """
    Handles both creating a new resume and editing an existing one.
    If resume_id is provided, it fetches the existing resume (draft or ready).
    Otherwise, it prepares for a new resume.
    """
    resume_instance = None
    if resume_id:
        resume_instance = get_object_or_404(Resume, id=resume_id, user=request.user)

    # Generate default resume name only if creating a new resume
    default_name = ""
    if not resume_instance:
        base_name = "My Resume"
        counter = 1
        default_name = base_name
        
        if request.user.is_authenticated:
            existing_names = list(request.user.resumes.values_list('name', flat=True))
            while default_name in existing_names:
                counter += 1
                default_name = f"{base_name} {counter}"
    
    # Hero section content
    hero_content = {
        'title': 'Resume Builder & Optimizer',
        'description': 'Create or optimize your ATS-friendly resume with our professional templates and AI-powered suggestions. Stand out to employers and increase your chances of getting hired.',
        'buttons': {
            'create': 'Create New Resume',
            'optimize': 'Update Resume'
        },
        'credits_text': ''
    }

    # Get month/year options for date dropdowns
    date_options = get_month_year_options()

    if resume_instance:
        selected_template = normalize_template_id(resume_instance.template_id)
    else:
        selected_template = normalize_template_id(request.GET.get('template'))

    resume_template_sections = get_resume_template_gallery_sections()
    selected_template_section_index = gallery_section_index_for_template_id(
        resume_template_sections, selected_template
    )

    profile_photo_preview_url = ''
    if resume_instance:
        profile_photo_preview_url = (
            augment_resume_dict_for_rendering(
                {
                    'personal_info': resume_instance.personal_info or {},
                    'experience': [],
                    'education': [],
                    'skills': {},
                    'additional': {},
                },
                request=request,
                resume_id=resume_instance.pk,
            )['personal_info'].get('profile_photo_display_url', '')
        )

    context = {
        'form': ResumeForm(),
        'hero_content': hero_content,
        'resume_services': {
            'builder_cost': 10  # or whatever the cost is
        },
        'default_resume_name': default_name,
        'resume_instance': resume_instance,
        'date_options': date_options,
        'resume_templates': templates_for_gallery(),
        'resume_template_sections': resume_template_sections,
        'selected_template': selected_template,
        'selected_template_section_index': selected_template_section_index,
        'default_template_id': DEFAULT_TEMPLATE_ID,
        'profile_photo_preview_url': profile_photo_preview_url,
        'selected_template_capabilities': template_ui_capabilities(selected_template),
    }

    return render(request, 'resume_builder/resume.html', context)

@login_required
def optimize_resume(request):
    if request.method == 'POST':
        try:
            # Debug logging
            logger.debug(f"Form data: {request.POST}")
            logger.debug(f"Files: {request.FILES}")

            resume_text = ""
            resume_file = request.FILES.get('resume_file')
            
            if resume_file:
                logger.info(f"Processing file: {resume_file.name}")
                # Get file extension
                file_extension = os.path.splitext(resume_file.name)[1].lower()
                
                # Create a temporary file to handle the upload
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    for chunk in resume_file.chunks():
                        temp_file.write(chunk)
                
                try:
                    # Process different file types
                    if file_extension == '.pdf':
                        logger.info("Processing PDF file")
                        with open(temp_file.name, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            for page in pdf_reader.pages:
                                resume_text += page.extract_text()
                    elif file_extension in ['.doc', '.docx']:
                        logger.info("Processing Word document")
                        doc = Document(temp_file.name)
                        resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        raise ValueError(f"Unsupported file type: {file_extension}")
                except Exception as e:
                    logger.error(f"File processing error: {str(e)}")
                    raise
                finally:
                    # Clean up the temporary file
                    os.unlink(temp_file.name)
            else:
                logger.info("Using text input")
                resume_text = request.POST.get('resume_text', '').strip()

            if not resume_text:
                logger.error("No resume content provided")
                raise ValueError("Please provide either a resume file or paste resume text")

            job_description = request.POST.get('job_description', '').strip()
            if not job_description:
                logger.error("No job description provided")
                raise ValueError("Please provide a job description")

            logger.info("Starting resume optimization")
            optimizer = ResumeOptimizer(api_key=os.getenv('OPENAI_API_KEY'))
            optimization_result = optimizer.optimize(resume_text, job_description)
            logger.info("Resume optimization completed")

            # Store optimization results in session
            request.session['optimization_results'] = {
                'original_content': resume_text,
                'optimized_content': optimization_result.optimized_content,
                'job_description': job_description,
                'keyword_matches': optimization_result.keyword_matches,
                'improvement_suggestions': optimization_result.improvement_suggestions,
                'ats_score': optimization_result.ats_score,
                'is_new': True  # Flag to indicate this is not saved yet
            }

            return JsonResponse({
                'success': True,
                'message': 'Resume optimized successfully',
                'redirect_url': reverse('resume_builder:view_resume')
            })

        except ValueError as ve:
            logger.warning(f"Validation error: {str(ve)}")
            return JsonResponse({
                'success': False,
                'error': str(ve)
            }, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}", exc_info=True)
            return JsonResponse({
                'success': False,
                'error': f"An error occurred: {str(e)}"
            }, status=400)
    
    # GET request
    hero_content = {
        'page_title': 'Optimize Your Resume',
        'page_description': 'Upload your resume and job description to create an ATS-optimized version tailored to the position.'
    }
    return render(request, 'resume_builder/optimize_resume.html', {'hero_content': hero_content})

@login_required
def upload_resume(request):
    """Upload and parse resume from file"""
    if request.method == 'POST':
        try:
            # Get the uploaded file
            resume_file = request.FILES.get('resume_file')
            template_id = normalize_template_id(request.POST.get('template_id'))
            
            if not resume_file:
                return JsonResponse({
                    'success': False,
                    'error': 'Please upload a resume file to continue. You can upload PDF, DOC, or DOCX files.'
                }, status=400)
            
            # Validate file type
            file_extension = os.path.splitext(resume_file.name)[1].lower()
            if file_extension not in ['.pdf', '.doc', '.docx']:
                return JsonResponse({
                    'success': False,
                    'error': 'Unsupported file format. Please upload a PDF, DOC, or DOCX file.'
                }, status=400)
            
            # Extract text from the file
            resume_text = ""
            
            # Create a temporary file to handle the upload
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                for chunk in resume_file.chunks():
                    temp_file.write(chunk)
                
                try:
                    # Process different file types
                    if file_extension == '.pdf':
                        logger.info("Processing PDF file")
                        try:
                            import pdfplumber
                            with pdfplumber.open(temp_file.name) as pdf:
                                for page in pdf.pages:
                                    resume_text += page.extract_text() or ""
                        except ImportError:
                            # Fallback to PyPDF2 if pdfplumber is not available
                            with open(temp_file.name, 'rb') as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                for page in pdf_reader.pages:
                                    resume_text += page.extract_text()
                    elif file_extension in ['.doc', '.docx']:
                        logger.info("Processing Word document")
                        doc = Document(temp_file.name)
                        resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except Exception as e:
                    logger.error(f"File processing error: {str(e)}")
                    return JsonResponse({
                        'success': False,
                        'error': f'Error processing file: {str(e)}'
                    }, status=400)
                finally:
                    # Clean up the temporary file
                    os.unlink(temp_file.name)
            
            if not resume_text.strip():
                return JsonResponse({
                    'success': False,
                    'error': 'Could not extract text from the uploaded file. Please ensure the file contains readable text.'
                }, status=400)
            
            # Parse the resume text using AI
            logger.info("Parsing resume with AI")
            try:
                parsed_data = format_resume(resume_text)
            except TimeoutError:
                logger.error("AI parsing timeout error")
                error_dialog = get_network_timeout_dialog()
                return JsonResponse({
                    'success': False,
                    'error': error_dialog
                }, status=504)
            except ConnectionError:
                logger.error("AI parsing connection error")
                error_dialog = get_network_connection_dialog()
                return JsonResponse({
                    'success': False,
                    'error': error_dialog
                }, status=503)
            except Exception as e:
                logger.error(f"AI parsing error: {str(e)}")
                # Check if it's a timeout-related error
                if "timed out" in str(e).lower() or "timeout" in str(e).lower():
                    error_dialog = get_network_timeout_dialog()
                    return JsonResponse({
                        'success': False,
                        'error': error_dialog
                    }, status=504)
                else:
                    error_dialog = get_network_generic_dialog()
                    return JsonResponse({
                        'success': False,
                        'error': error_dialog
                    }, status=500)
            
            personal_info = parsed_data.personal_info.model_dump()
            experience = parsed_data.experience
            skills = parsed_data.skills.model_dump()
            
            # Always save as draft for user verification
            should_be_draft = True
            
            # Create the resume object
            resume_name = personal_info.get('full_name', 'Uploaded Resume')
            if not resume_name or resume_name == 'Uploaded Resume':
                resume_name = f"Resume from {resume_file.name}"
            
            resume = Resume.objects.create(
                user=request.user,
                name=resume_name,
                draft=should_be_draft,
                template_id=template_id,
                original_content=resume_text,
                personal_info=personal_info,
                experience=[exp.model_dump() for exp in experience],
                education=[edu.model_dump() for edu in parsed_data.education],
                skills=skills,
                additional=parsed_data.additional.model_dump(),
                is_optimized=False,  # This is not an optimized resume
                # Set default values for optimization fields
                keyword_matches=[],
                improvement_suggestions=[],
                ats_score=0
            )
            
            logger.info(f"Resume created successfully. ID: {resume.id}, Draft: {should_be_draft}")
            
            # Always redirect to edit page for verification
            redirect_url = reverse('resume_builder:edit_resume', args=[resume.id])
            message = 'Resume uploaded and parsed. Please review and verify all information before marking as ready.'
            
            return JsonResponse({
                'success': True,
                'message': message,
                'resume_id': resume.id,
                'redirect_url': redirect_url,
                'is_draft': should_be_draft
            })
            
        except Exception as e:
            logger.error(f"Unexpected error in upload_resume: {str(e)}", exc_info=True)
            error_dialog = get_network_generic_dialog()
            return JsonResponse({
                'success': False,
                'error': error_dialog
            }, status=500)
    
    # GET request - render the upload form
    hero_content = {
        'page_title': 'Upload Resume',
        'page_description': 'Upload your existing resume to edit or convert it to a new format.'
    }
    return render(
        request,
        'resume_builder/optimize_resume_form.html',
        {
            'hero_content': hero_content,
            'resume_templates': templates_for_gallery(),
            'resume_template_sections': get_resume_template_gallery_sections(),
            'selected_template': DEFAULT_TEMPLATE_ID,
            'default_template_id': DEFAULT_TEMPLATE_ID,
        },
    )

@login_required
def view_resume(request, resume_id=None):
    """
    Display a resume by rendering the template dynamically from database data.
    If resume_id is provided, it fetches a saved resume and renders it with the chosen template.
    Otherwise, it uses unsaved optimization results from the session.
    
    Supports HTMX requests for AI assistant integration by returning only the resume content.
    Supports template switching via query parameter 'template'.
    """
    context = {}
    if resume_id:
        # Viewing a saved resume
        resume = get_object_or_404(Resume, id=resume_id, user=request.user)
        
        # Prepare resume data for template rendering
        resume_data = {
            'personal_info': resume.personal_info or {},
            'experience': resume.experience or [],
            'education': resume.education or [],
            'skills': resume.skills or {},
            'additional': resume.additional or {}
        }
        
        # Get the template ID - support template switching via query parameter
        template_id = normalize_template_id(
            request.GET.get('template') or resume.template_id or None
        )
        
        # Render the resume with the chosen template
        html_content = _render_resume_template_html(
            request, resume_data, template_id, resume_id=resume.pk
        )
        
        # Check if this is an HTMX request (for AI assistant integration)
        is_htmx_request = request.headers.get('HX-Request') == 'true'
        
        if is_htmx_request:
            # For HTMX requests, return only the resume content without full page layout
            # Update the resume's template_id in the context to match the selected template
            resume.template_id = template_id
            context = {
                'resume': resume,
                'resume_html': html_content,
                'is_htmx_request': True,
                'resume_templates': templates_for_gallery(),
                'resume_template_sections': get_resume_template_gallery_sections(),
                'default_template_id': DEFAULT_TEMPLATE_ID,
            }
            return render(request, 'resume_builder/component/resume_preview_tab.html', context)
        
        # Create full HTML document for regular requests
        full_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume.personal_info.get('full_name', 'Resume')} - Resume</title>
    <script src="https://cdn.tailwindcss.com"></script>
{get_resume_embedded_style_tag()}
</head>
<body class="bg-white">
    {html_content}
</body>
</html>
"""
        
        context = {
            'resume': resume,
            'resume_html': html_content,
            'hero_content': {
                'page_title': 'View Resume',
                'page_description': 'Review your generated resume below. You can edit it further or download it.'
            },
            'default_template_id': DEFAULT_TEMPLATE_ID,
        }
        
        return render(request, 'resume_builder/view_resume.html', context)
    else:
        # Viewing unsaved optimization results from session
        optimization_results = request.session.get('optimization_results')
        
        if not optimization_results:
            # If no session data, redirect to the start
            return redirect('resume_builder:optimize_resume')
        
        # Prepare data for template rendering (simplified version)
        resume_data = {
            'optimized_content': optimization_results.get('optimized_content', '')
        }
        
        context = {
            'resume': None, # No saved resume object
            'resume_html': None, # No HTML for this case
            'resume_data': resume_data,
            'hero_content': {
                'page_title': 'Optimized Resume',
                'page_description': 'Review your optimized resume content. You can now save or download it.'
            },
            'default_template_id': DEFAULT_TEMPLATE_ID,
        }
        return render(request, 'resume_builder/view_resume.html', context)

@login_required
def download_resume(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content', '')

            # Create a new Word document
            doc = Document()
            doc.add_heading('Optimized Resume', 0)

            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # Save to BytesIO instead of temporary file
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # For direct download
            response = HttpResponse(
                doc_io.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=optimized_resume.docx'
            return response

        except Exception as e:
            logger.error(f"Word generation error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)

    # GET request - Get resume data from session
    resume_data = request.session.get('resume_data')
    if not resume_data:
        return redirect('resume_builder:create_resume')
    
    templates = templates_for_download_picker()

    active_template = normalize_template_id(request.session.get('active_template'))
    
    context = {
        'resume_data': resume_data,
        'templates': templates,
        'active_template': active_template,
        'selected_template': f'resume_templates/{active_template}.html'
    }
    
    return render(request, 'download_resume.html', context)

@login_required
def switch_template(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            template_id = normalize_template_id(data.get('template_id'))
            
            # Store selected template in session
            request.session['active_template'] = template_id
            
            # Render the selected template with resume data
            resume_data = request.session.get('resume_data') or {}
            augmented = augment_resume_dict_for_rendering(resume_data, request=request)
            return render(request, f'resume_templates/{template_id}.html', {'resume_data': augmented})
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

def preview_template(request, template_id):
    """Preview template with sample data - optimized for multi-country/multi-language support"""
    if not is_valid_template_id(template_id):
        return JsonResponse({'error': 'Invalid template ID'}, status=400)
    
    # Get user's locale from request (could be from user preferences, Accept-Language header, etc.)
    user_locale = request.GET.get('locale', 'en-US')  # Default to US English
    
    # Sample resume data - optimized for internationalization
    # In production, this would come from a database or translation files
    sample_resume_data = get_localized_sample_data(user_locale)
    augmented = augment_resume_dict_for_rendering(sample_resume_data, request=request)

    context = {'resume_data': augmented}
    return render(request, f'resume_templates/{template_id}.html', context)

def get_localized_sample_data(locale='en-US'):
    """Get sample resume data based on locale - supports multi-country/multi-language"""
    
    # Sample data templates for different locales
    sample_data_templates = {
        'en-US': {
            'personal_info': {
                'full_name': 'Sarah Johnson',
                'title': 'Senior Software Engineer',
                'email': 'sarah.johnson@email.com',
                'phone': '(555) 123-4567',
                'location': 'Austin, TX',
                'street_address': '',
                'linkedin': 'https://www.linkedin.com/in/sarahjohnson',
                'profile_photo_display_url': 'https://images.unsplash.com/photo-1494790108377-be9c29b29330?auto=format&fit=crop&w=320&q=72',
                'summary': 'Experienced software engineer with 5+ years developing scalable web applications. Passionate about clean code, user experience, and emerging technologies.'
            },
            'experience': [
                {
                    'company': 'TechCorp Inc.',
                    'title': 'Senior Software Engineer',
                    'position': 'Senior Software Engineer',
                    'start_date': '2022-01',
                    'startDate': '2022-01',
                    'end_date': 'Present',
                    'endDate': 'Present',
                    'description': '<ul><li>Led microservices architecture for 1M+ users.</li><li>Mentored engineers and rolled out CI/CD.</li></ul>',
                },
                {
                    'company': 'StartupXYZ',
                    'title': 'Full Stack Developer',
                    'position': 'Full Stack Developer',
                    'start_date': '2020-03',
                    'startDate': '2020-03',
                    'end_date': '2021-12',
                    'endDate': '2021-12',
                    'description': '<ul><li>Shipped React/Node apps end-to-end.</li><li>Partnered with design on responsive UI.</li></ul>',
                },
            ],
            'education': [
                {
                    'institution': 'University of Technology',
                    'degree': 'Bachelor of Computer Science',
                    'start_date': '2016-09',
                    'startDate': '2016-09',
                    'end_date': '2020-05',
                    'endDate': '2020-05',
                    'description': 'Graduated with honors. Specialized in software engineering and database systems.',
                }
            ],
            'skills': {
                'technical': ['JavaScript', 'React', 'Node.js', 'Python', 'PostgreSQL', 'AWS'],
                'soft': ['Leadership', 'Problem Solving', 'Team Collaboration', 'Communication'],
                'languages': ['English', 'Spanish'],
                'rated': [
                    {'name': 'JavaScript', 'level': 9},
                    {'name': 'React', 'level': 8},
                    {'name': 'Python', 'level': 8},
                    {'name': 'System design', 'level': 7},
                ],
            },
            'additional': {
                'certifications': '<ul><li>AWS Certified Developer</li><li>Google Cloud Professional</li></ul>',
                'projects': '<ul><li>Open-source contributor to the React ecosystem</li><li>Personal finance tracker (React, Node, PostgreSQL)</li></ul>',
                'references': [
                    {
                        'name': 'Alex Morgan',
                        'affiliation': 'Engineering Director · TechCorp Inc.',
                        'email': 'alex.morgan@techcorp.example',
                        'phone': '(555) 010-2030',
                    },
                    {
                        'name': 'Jordan Smith',
                        'affiliation': 'CTO · StartupXYZ',
                        'email': 'jsmith@startupxyz.example',
                        'phone': '',
                    },
                ],
            }
        },
        'es-ES': {
            'personal_info': {
                'full_name': 'María García López',
                'title': 'Ingeniera de Software Senior',
                'email': 'maria.garcia@email.com',
                'phone': '+34 612 345 678',
                'summary': 'Ingeniera de software experimentada con más de 5 años desarrollando aplicaciones web escalables. Apasionada por el código limpio, la experiencia de usuario y las tecnologías emergentes.'
            },
            'experience': [
                {
                    'company': 'TechCorp España',
                    'title': 'Ingeniera de Software Senior',
                    'position': 'Ingeniera de Software Senior',
                    'start_date': '2022-01',
                    'startDate': '2022-01',
                    'end_date': 'Presente',
                    'endDate': 'Presente',
                    'description': '<ul><li>Arquitectura de microservicios para más de 1M de usuarios.</li><li>Mentoría e implementación de CI/CD.</li></ul>',
                }
            ],
            'education': [
                {
                    'institution': 'Universidad Politécnica de Madrid',
                    'degree': 'Grado en Ingeniería Informática',
                    'start_date': '2016-09',
                    'startDate': '2016-09',
                    'end_date': '2020-05',
                    'endDate': '2020-05',
                    'description': 'Graduada con honores. Especializada en ingeniería de software y sistemas de bases de datos.',
                }
            ],
            'skills': {
                'technical': ['JavaScript', 'React', 'Node.js', 'Python', 'PostgreSQL', 'AWS'],
                'soft': ['Liderazgo', 'Resolución de Problemas', 'Colaboración en Equipo', 'Comunicación'],
                'languages': ['Español', 'Inglés']
            },
            'additional': {
                'certifications': 'Desarrollador Certificado AWS, Profesional de Google Cloud',
                'projects': 'Contribuidora de código abierto al ecosystème React, Construí aplicación de seguimiento financiero personal'
            }
        },
        'fr-FR': {
            'personal_info': {
                'full_name': 'Sophie Martin',
                'title': 'Ingénieure Logiciel Senior',
                'email': 'sophie.martin@email.com',
                'phone': '+33 1 23 45 67 89',
                'summary': 'Ingénieure logiciel expérimentée avec plus de 5 ans de développement d\'applications web évolutives. Passionnée par le code propre, l\'expérience utilisateur et les technologies émergentes.'
            },
            'experience': [
                {
                    'company': 'TechCorp France',
                    'title': 'Ingénieure Logiciel Senior',
                    'position': 'Ingénieure Logiciel Senior',
                    'start_date': '2022-01',
                    'startDate': '2022-01',
                    'end_date': 'Présent',
                    'endDate': 'Présent',
                    'description': '<ul><li>Architecture microservices desservant plus d\'1M d\'utilisateurs.</li><li>Encadrement de développeurs juniors et pipelines CI/CD.</li></ul>',
                }
            ],
            'education': [
                {
                    'institution': 'École Centrale Paris',
                    'degree': 'Master en Informatique',
                    'start_date': '2016-09',
                    'startDate': '2016-09',
                    'end_date': '2020-05',
                    'endDate': '2020-05',
                    'description': 'Diplômée avec mention. Spécialisée en génie logiciel et systèmes de bases de données.',
                }
            ],
            'skills': {
                'technical': ['JavaScript', 'React', 'Node.js', 'Python', 'PostgreSQL', 'AWS'],
                'soft': ['Leadership', 'Résolution de Problèmes', 'Collaboration d\'Équipe', 'Communication'],
                'languages': ['Français', 'Anglais']
            },
            'additional': {
                'certifications': 'Développeur Certifié AWS, Professionnel Google Cloud',
                'projects': 'Contributeur open source à l\'écosystème React, Construit une application de suivi financier personnel'
            }
        }
    }
    
    # Return localized data or fallback to US English
    return sample_data_templates.get(locale, sample_data_templates['en-US'])

def resume_home(request):
    return render(request, 'resume_builder/resume.html')

@login_required
def save_resume(request):
    if request.method == 'POST':
        try:
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                resume_id = data.get('resume_id')

                if resume_id:
                    # Find existing resume and update it
                    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
                    message = 'Changes saved successfully!'
                else:
                    # Create a new resume instance
                    resume = Resume(user=request.user)
                    message = 'Optimized resume saved successfully!'

                # Update the resume object with data from the request
                resume.optimized_content = data.get('content', '')
                resume.job_description = data.get('job_description', '')
                resume.ats_score = data.get('ats_score', 0)
                resume.keyword_matches = data.get('keyword_matches', [])
                resume.improvement_suggestions = data.get('improvement_suggestions', [])
                resume.template_id = normalize_template_id(data.get('template_id'))
                
                # After updating content, we must regenerate the HTML file
                from django.template.loader import render_to_string
                from django.core.files.base import ContentFile
                
                resume_template_data = {
                    'personal_info': {
                        'full_name': f'Resume #{resume.id}' if resume.id else 'Optimized Resume',
                        'title': 'ATS-Optimized Version',
                        'email': '', 'phone': '',
                        'summary': resume.optimized_content
                    },
                    'experience': [], 'education': [],
                    'skills': {'technical': resume.keyword_matches, 'soft': [], 'languages': []},
                    'additional': {
                        'certifications': f'ATS Score: {resume.ats_score}/100',
                        'projects': 'Optimized for Applicant Tracking Systems'
                    }
                }
                
                html_content = _render_resume_template_html(
                    request,
                    resume_template_data,
                    resume.template_id,
                    resume_id=resume.pk if resume.pk else None,
                )
                
                full_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Optimized Resume</title>
    <script src="https://cdn.tailwindcss.com"></script>
{get_resume_embedded_style_tag()}
</head>
<body class="bg-white">
    {html_content}
</body>
</html>
"""
                file_content = ContentFile(full_html.encode('utf-8'))
                
                # Save the resume object first to ensure it has an ID
                if not resume.id:
                    resume.save()

                # Now save the file with a name that includes the ID
                resume.pdf_file.save(f'optimized_resume_{resume.id}.html', file_content, save=True)

                return JsonResponse({
                    'success': True,
                    'message': message,
                    'resume_id': resume.id
                })
            else:
                # Handle PDF file upload (existing functionality)
                pdf_file = request.FILES.get('pdf_file')
                template_id = normalize_template_id(request.POST.get('template_id'))
                
                if not pdf_file:
                    return JsonResponse({'success': False, 'error': 'PDF file is required'})
                
                # Save the PDF file with default values for required fields
                resume = Resume.objects.create(
                    user=request.user,
                    pdf_file=pdf_file,
                    template_id=template_id,
                    original_content='',
                    optimized_content='',
                    job_description='',
                    ats_score=0,
                    keyword_matches=[],
                    improvement_suggestions=[]
                )
                
                return JsonResponse({
                    'success': True,
                    'message': 'Resume PDF saved successfully!',
                    'resume_id': resume.id
                })
            
        except Exception as e:
            logger.error(f"Save resume error: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)})
            
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@login_required
def create_resume_submit(request):
    if request.method == 'POST':
        try:
            # Parse the JSON data from request
            data = json.loads(request.body)
            
            # Store resume data in session for template rendering
            request.session['resume_data'] = {
                'personal_info': {
                    'full_name': data.get('fullName'),
                    'title': data.get('title'),
                    'email': data.get('email'),
                    'phone': data.get('phone'),
                    'summary': data.get('summary')
                },
                'experience': data.get('experience', []),
                'education': data.get('education', []),
                'skills': {
                    'technical': data.get('technicalSkills', '').split(','),
                    'soft': data.get('softSkills', '').split(','),
                    'languages': data.get('languages', '').split(',')
                },
                'additional': {
                    'certifications': data.get('certifications'),
                    'projects': data.get('projects')
                }
            }

            return JsonResponse({
                'success': True,
                'message': 'Resume created successfully',
                'redirect_url': reverse('resume_builder:download_resume')
            })

        except json.JSONDecodeError as e:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            logger.error(f"Resume creation error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to create resume'
            }, status=500)

    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)

@login_required
def my_resumes(request):
    """List all resumes for the current user, with ready resumes first."""
    resumes = request.user.resumes.filter(is_optimized=False).order_by('draft', '-updated_at')
    
    context = {
        'resumes': resumes,
        'hero_content': {
            'page_title': 'My Resumes',
            'page_description': 'Manage and view all your saved resumes.'
        }
    }
    
    return render(request, 'resume_builder/my_resumes.html', context)

@login_required
def get_resume_content(request, resume_id):
    """
    Fetch and render only the HTML content of a specific resume.
    This is used to populate the preview modal without a full page reload.
    """
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    resume_data = {
        'personal_info': resume.personal_info or {},
        'experience': resume.experience or [],
        'education': resume.education or [],
        'skills': resume.skills or {},
        'additional': resume.additional or {}
    }
    
    template_id = normalize_template_id(resume.template_id)
    html_content = _render_resume_template_html(
        request, resume_data, template_id, resume_id=resume.pk
    )
    
    return HttpResponse(html_content)

@login_required
def delete_resume(request, resume_id):
    """Delete a resume and its associated file"""
    if request.method == 'POST':
        try:
            # Get the resume and ensure it belongs to the current user
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            
            # Delete the resume (this will also delete the file due to the model's delete method)
            resume.delete()
            
            return JsonResponse({
                'success': True,
                'message': 'Resume deleted successfully!'
            })
            
        except Resume.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Resume not found'
            }, status=404)
        except Exception as e:
            logger.error(f"Delete resume error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to delete resume'
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)

@csrf_exempt
def download_resume_file(request, resume_id, format_type='html'):
    """Download a resume in the specified format - supports both authenticated and anonymous users"""
    try:
        if str(resume_id) == 'anonymous':
            # Anonymous download: expect POSTed JSON data
            if request.method != 'POST':
                return JsonResponse({'error': 'Anonymous download requires POST'}, status=405)
            try:
                data = json.loads(request.body)
            except Exception as e:
                return JsonResponse({'error': 'Invalid JSON data'}, status=400)
            resume_data, template_id = resume_data_and_template_from_wizard_payload(
                data,
                fallback_template_id=DEFAULT_TEMPLATE_ID,
            )
        else:
            # Authenticated user: GET uses DB snapshot; POST uses same JSON shape as anonymous
            # (wizard draft preview/download with optional profile_photo_display_url data URL).
            if not request.user.is_authenticated:
                return JsonResponse({'error': 'Authentication required'}, status=403)

            resume = get_object_or_404(Resume, id=resume_id, user=request.user)

            snapshot_from_body = request.method == 'POST' and bool(request.body)

            if snapshot_from_body:
                try:
                    data = json.loads(request.body.decode('utf-8'))
                except (json.JSONDecodeError, UnicodeDecodeError):
                    return JsonResponse({'error': 'Invalid JSON data'}, status=400)
                body_rid = data.get('resume_id')
                if body_rid is None or str(body_rid) != str(resume_id):
                    return JsonResponse({'error': 'resume_id required and must match URL'}, status=400)
                resume_data, template_id = resume_data_and_template_from_wizard_payload(
                    data,
                    fallback_template_id=resume.template_id or DEFAULT_TEMPLATE_ID,
                )
            else:
                if request.method != 'GET':
                    return JsonResponse({'error': 'Method not allowed'}, status=405)
                resume_data = {
                    'personal_info': resume.personal_info or {},
                    'experience': resume.experience or [],
                    'education': resume.education or [],
                    'skills': resume.skills or {},
                    'additional': resume.additional or {}
                }
                template_id = normalize_template_id(resume.template_id)

        augmented_resume_data = augment_resume_dict_for_rendering(
            resume_data,
            request=request,
            resume_id=resume.pk,
            force_inline_profile_photo=True,
        )

        inline = (request.GET.get('inline') or '').strip().lower() in ('1', 'true', 'yes', 'on')

        # Now render and return in the requested format (reuse existing logic)
        if format_type == 'html':
            html_content = render_to_string(
                f'resume_templates/{template_id}.html',
                {'resume_data': augmented_resume_data},
            )
            full_html = f"""
<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <title>{augmented_resume_data['personal_info'].get('full_name', 'Resume')} - Resume</title>
    <script src=\"https://cdn.tailwindcss.com\"></script>
{get_resume_embedded_style_tag()}
</head>
<body class=\"bg-white\">
    {html_content}
</body>
</html>
"""
            response = HttpResponse(full_html, content_type='text/html')
            if inline:
                response['Content-Disposition'] = (
                    f'inline; filename="resume_{resume_id}_draft.html"'
                )
            else:
                response['Content-Disposition'] = (
                    f'attachment; filename="resume_{resume_id}.html"'
                )
            return response
        elif format_type == 'pdf':
            try:
                from pdf_generator.core.generator import PDFGenerator

                context = {
                    'resume_data': augmented_resume_data,
                    'resume_name': augmented_resume_data['personal_info'].get('full_name', 'Resume'),
                    'generated_date': timezone.now().strftime('%B %d, %Y'),
                }
                pdf_bytes = PDFGenerator.generate_from_template(
                    template_name=f'resume_templates/{template_id}.html',
                    context=context,
                    options={
                        'format': 'A4',
                        'print_background': True,
                        'margins': {
                            'top': '0.5in',
                            'right': '0.5in',
                            'bottom': '0.5in',
                            'left': '0.5in',
                        },
                    },
                )
                response = HttpResponse(pdf_bytes, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="resume_{resume_id}.pdf"'
                return response
            except ImportError:
                logger.warning(
                    "pdf_generator package not available for resume PDF download."
                )
                return JsonResponse(
                    {'error': 'PDF generation package is not available'}, status=500
                )
        elif format_type == 'word':
            try:
                from html2docx import html2docx
                html_content = render_to_string(
                    f'resume_templates/{template_id}.html',
                    {'resume_data': augmented_resume_data},
                )
                full_html = f"""
<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <title>{augmented_resume_data['personal_info'].get('full_name', 'Resume')} - Resume</title>
    <script src=\"https://cdn.tailwindcss.com\"></script>
</head>
<body class=\"bg-white\">
    {html_content}
</body>
</html>
"""
                docx_buffer = html2docx(full_html, title=f"Resume - {augmented_resume_data['personal_info'].get('full_name', 'Resume')}")
                response = HttpResponse(
                    docx_buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="resume_{resume_id}.docx"'
                return response
            except ImportError:
                logger.error("html2docx not available for Word document generation")
                return JsonResponse({'error': 'Word document generation not available'}, status=500)
        else:
            return JsonResponse({'error': 'Invalid format'}, status=400)
    except Exception as e:
        logger.error(f"Download resume error: {str(e)}")
        return JsonResponse({'error': 'Failed to download resume'}, status=500)

@login_required
def check_resume_update_access(request):
    """Check if user has access to resume update functionality"""
    access_info = check_subscription_access(
        user=request.user,
        # This flow maps to the saved resume capability in the seeded feature catalog.
        feature_identifier='resume_saving',
        required_plan='Plus'
    )
    
    return JsonResponse(access_info)

@login_required
def save_personal_info(request):
    """Save step 1: Personal Information and create or update resume"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)

            if resume_id:
                # Update existing resume
                resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            else:
                # Create new draft resume
                resume = Resume.objects.create(
                    user=request.user,
                    name=data.get('resume_name', 'My Resume'),
                    draft=True,
                    template_id=normalize_template_id(data.get('template_id')),
                )
            # Defensive update: only update summary if present in payload
            personal_info = resume.personal_info or {}
            personal_info.pop('profile_photo_display_url', None)
            personal_info['full_name'] = data.get('fullName', personal_info.get('full_name', ''))
            personal_info['title'] = data.get('title', personal_info.get('title', ''))
            personal_info['email'] = data.get('email', personal_info.get('email', ''))
            personal_info['phone'] = data.get('phone', personal_info.get('phone', ''))
            if 'location' in data:
                personal_info['location'] = data.get('location', personal_info.get('location', ''))
            if 'street_address' in data:
                personal_info['street_address'] = data.get(
                    'street_address', personal_info.get('street_address', '')
                )
            if 'linkedin' in data:
                personal_info['linkedin'] = data.get('linkedin', personal_info.get('linkedin', ''))
            if 'summary' in data:
                personal_info['summary'] = data.get('summary', personal_info.get('summary', ''))
            resume.name = data.get('resume_name', resume.name)
            resume.personal_info = personal_info
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            resume.save()
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Personal information saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })
            return JsonResponse({
                'success': True,
                'message': 'Personal information saved successfully!',
                'resume_id': resume.id
            })
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save personal info error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save personal information'}, status=500)
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def save_experience(request):
    """Save step 2: Work Experience"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)
            
            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            resume.experience = data.get('experience', [])
            
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            
            resume.save()
            
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Experience saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })

            return JsonResponse({'success': True, 'message': 'Experience saved successfully!'})
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save experience error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save experience'}, status=500)
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def save_education(request):
    """Save step 3: Education"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)

            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            resume.education = data.get('education', [])
            
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            
            resume.save()
            
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Education saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })

            return JsonResponse({'success': True, 'message': 'Education saved successfully!'})
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save education error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save education'}, status=500)
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def save_skills(request):
    """Save step 4: Skills"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)

            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            prev_skills = resume.skills or {}
            rated_kw = None
            if 'rated_skills' in data:
                rated_kw = data.get('rated_skills')
            elif 'rated' in data:
                rated_kw = data.get('rated')
            resume.skills = merge_skills_payload(
                prev_skills,
                technical=[
                    skill.strip()
                    for skill in data.get('technicalSkills', '').split(',')
                    if skill.strip()
                ],
                soft=[
                    skill.strip()
                    for skill in data.get('softSkills', '').split(',')
                    if skill.strip()
                ],
                languages=[
                    lang.strip()
                    for lang in data.get('languages', '').split(',')
                    if lang.strip()
                ],
                rated_raw=rated_kw if ('rated_skills' in data or 'rated' in data) else ...,
            )
            
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            
            resume.save()
            
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Skills saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })

            return JsonResponse({'success': True, 'message': 'Skills saved successfully!'})
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save skills error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save skills'}, status=500)
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def save_additional(request):
    """Save step 5: Additional Information"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)
            
            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            resume.additional = merge_additional_payload(
                resume.additional,
                {
                    'certifications': data.get('certifications'),
                    'projects': data.get('projects'),
                    **({'references': data['references']} if 'references' in data else {}),
                },
            )
            
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            
            resume.save()
            
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Additional info saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })

            return JsonResponse({'success': True, 'message': 'Additional information saved successfully!'})
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save additional error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save additional information'}, status=500)
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
@require_GET
def profile_photo_proxy(request, resume_id):
    """
    Serve stored portrait bytes for the authenticated owner (same-origin <img src>).
    Avoids V4 signed URLs, which require a private key — unavailable on Cloud Run ADC.
    """
    resume = get_object_or_404(Resume, pk=resume_id, user=request.user)
    pi = resume.personal_info or {}
    bucket_name = pi.get("profile_image_gcs_bucket")
    blob_name = pi.get("profile_image_blob")

    if (
        bucket_name
        and blob_name
        and getattr(settings, "ENABLE_GCS_PROFILE_UPLOAD", False)
    ):
        try:
            from google.cloud import storage

            client = storage.Client(
                project=(getattr(settings, "GCP_PROJECT_ID", "") or None)
            )
            blob = client.bucket(bucket_name).blob(blob_name)
            data = blob.download_as_bytes()
            if not data:
                raise Http404("Portrait object empty")
            ctype = (
                (blob.content_type or "").split(";")[0].strip()
                or mimetypes.guess_type(blob_name)[0]
                or "application/octet-stream"
            )
            resp = HttpResponse(data, content_type=ctype)
            resp["Cache-Control"] = "private, max-age=3600"
            return resp
        except Http404:
            raise
        except Exception:
            logger.exception("profile_photo_proxy GCS read failed resume_id=%s", resume_id)
            raise Http404("Portrait not available")

    local_path = pi.get("profile_image_local_path")
    if local_path:
        from django.core.files.storage import default_storage

        rel = str(local_path).lstrip("/").replace("\\", "/")
        if default_storage.exists(rel):
            with default_storage.open(rel, "rb") as fh:
                data = fh.read()
            ctype = mimetypes.guess_type(rel)[0] or "application/octet-stream"
            resp = HttpResponse(data, content_type=ctype)
            resp["Cache-Control"] = "private, max-age=3600"
            return resp

    raise Http404("No portrait for this resume")

@login_required
@require_http_methods(["POST"])
def upload_profile_photo(request):
    """Multipart upload for optional profile portrait (stored in GCS or local MEDIA)."""
    from .profile_photo_media import (
        ingest_profile_photo,
        profile_photo_storage_backend,
        resolve_profile_photo_display_url,
    )

    resume_id = request.POST.get('resume_id')
    if not resume_id:
        return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
    resume = get_object_or_404(Resume, pk=resume_id, user=request.user)
    upload = request.FILES.get('photo')
    if not upload:
        return JsonResponse(
            {'success': False, 'error': 'No photo file was received. Pick an image and try again.'},
            status=400,
        )
    if getattr(upload, 'size', None) == 0:
        return JsonResponse(
            {'success': False, 'error': 'The selected file was empty. Choose another image.'},
            status=400,
        )
    try:
        ingest_profile_photo(resume, upload)
    except ValueError as exc:
        return JsonResponse({'success': False, 'error': str(exc)}, status=400)
    except Exception:
        logger.exception('upload_profile_photo failed')
        return JsonResponse({'success': False, 'error': 'Failed to upload photo'}, status=500)

    resume.refresh_from_db()
    display_url = resolve_profile_photo_display_url(
        resume.personal_info or {}, request=request, resume_id=resume.pk
    )
    stored_as = profile_photo_storage_backend(resume.personal_info or {})
    payload = {'success': True, 'profile_photo_display_url': display_url, 'stored_as': stored_as}
    return JsonResponse(payload)


@login_required
@require_http_methods(["POST"])
def delete_profile_photo(request):
    """Remove stored portrait for this resume (GCS or local); clears JSON pointers."""
    from .profile_photo_media import clear_resume_profile_photo

    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    resume_id = data.get('resume_id')
    if not resume_id:
        return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
    resume = get_object_or_404(Resume, pk=resume_id, user=request.user)
    try:
        clear_resume_profile_photo(resume)
    except Exception:
        logger.exception('delete_profile_photo failed')
        return JsonResponse({'success': False, 'error': 'Failed to remove photo'}, status=500)
    return JsonResponse({'success': True})

@login_required
def finalize_resume(request):
    """Finalize resume with template - no longer generates HTML files"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            action = data.get('action', 'save')  # 'save' or 'preview'
            
            if not resume_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Resume ID is required'
                }, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            
            # Update template
            template_id = normalize_template_id(data.get('template_id'))
            resume.template_id = template_id
            
            # Set draft to False
            resume.draft = False
            resume.save()
            
            if action == 'save':
                return JsonResponse({
                    'success': True,
                    'message': 'Resume saved successfully!',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })
            else:  # action == 'preview'
                return JsonResponse({
                    'success': True,
                    'message': 'Resume created successfully',
                    'redirect_url': reverse('resume_builder:view_resume_by_id', args=[resume.id])
                })
            
        except json.JSONDecodeError as e:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            logger.error(f"Finalize resume error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to finalize resume'
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)

@csrf_exempt
def preview_anonymous_resume(request):
    """
    Accepts POSTed JSON data for an anonymous resume preview, renders the selected template,
    and returns the view_resume.html page with the rendered resume.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_data, template_id = resume_data_and_template_from_wizard_payload(
                data,
                fallback_template_id=DEFAULT_TEMPLATE_ID,
            )

            html_content = _render_resume_template_html(
                request,
                resume_data,
                template_id,
                resume_id=None,
                force_inline_profile_photo=False,
            )
            augmented = augment_resume_dict_for_rendering(
                resume_data,
                request=request,
                resume_id=None,
                force_inline_profile_photo=False,
            )

            context = {
                'resume': None,  # No saved resume object
                'resume_html': html_content,
                'resume_data': augmented,
                'hero_content': {
                    'page_title': 'Preview Resume',
                    'page_description': 'Review your generated resume below. You can download it or sign up to save it.'
                },
                'default_template_id': DEFAULT_TEMPLATE_ID,
            }
            return render(request, 'resume_builder/view_resume.html', context)
        except Exception as e:
            logger.error(f"Anonymous preview error: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def create_resume_from_data(request):
    """Create a resume from JSON data for users returning from authentication"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            # Extract resume data from the request and map field names
            raw_personal_info = data.get('personalInfo', {})
            display_url = (
                raw_personal_info.get('profile_photo_display_url')
                or raw_personal_info.get('profilePhotoDisplayUrl')
                or ''
            )
            personal_info = {
                'full_name': raw_personal_info.get('fullName', ''),
                'title': raw_personal_info.get('title', ''),
                'email': raw_personal_info.get('email', ''),
                'phone': raw_personal_info.get('phone', ''),
                'summary': raw_personal_info.get('summary', ''),
                'location': raw_personal_info.get('location', ''),
                'street_address': raw_personal_info.get('street_address', ''),
                'linkedin': raw_personal_info.get('linkedin', ''),
            }
            if display_url:
                personal_info['profile_photo_display_url'] = display_url
            
            experience = data.get('experience', [])
            education = data.get('education', [])
            skills = data.get('skills', {})
            additional = data.get('additional', {})
            template_id = normalize_template_id(data.get('templateId'))
            resume_name = data.get('resume_name', 'My Resume')
            
            # Create the resume object
            resume = Resume.objects.create(
                user=request.user,
                name=resume_name,
                draft=False,  # Mark as complete since user finished it
                template_id=template_id,
                personal_info=personal_info,
                experience=experience,
                education=education,
                skills=skills,
                additional=additional
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Resume created successfully from your saved progress!',
                'resume_id': resume.id,
                'redirect_url': reverse('resume_builder:my_resumes')
            })
            
        except json.JSONDecodeError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            logger.error(f"Error creating resume from data: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to create resume'
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)

@login_required
def create_resume_after_auth(request):
    """Handle resume creation after authentication and redirect to my-resumes"""
    if request.method == 'GET':
        # This view will be called after authentication
        # The actual resume creation will be handled by JavaScript
        return render(request, 'resume_builder/create_resume_after_auth.html')
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@login_required
def save_summary(request):
    """Save step 6: Professional Summary"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            is_editing_save = data.get('is_editing_save', False)
            
            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            # Only update the summary field in personal_info
            personal_info = resume.personal_info or {}
            personal_info['summary'] = data.get('summary', '')
            resume.personal_info = personal_info
            # Update template if provided
            if data.get('template_id'):
                resume.template_id = normalize_template_id(data.get('template_id'))
            resume.save()
            if is_editing_save:
                return JsonResponse({
                    'success': True,
                    'message': 'Summary saved.',
                    'redirect_url': reverse('resume_builder:my_resumes')
                })
            return JsonResponse({'success': True, 'message': 'Summary saved successfully!'})
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
        except Exception as e:
            logger.error(f"Save summary error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to save summary'}, status=500)
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

@login_required
def generate_ai_summary(request):
    """Generate a professional summary for a resume using AI."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume_id = data.get('resume_id')
            if not resume_id:
                return JsonResponse({'success': False, 'error': 'Resume ID is required'}, status=400)
            resume = get_object_or_404(Resume, id=resume_id, user=request.user)
            resume_data = {
                'personal_info': resume.personal_info or {},
                'experience': resume.experience or [],
                'education': resume.education or [],
                'skills': resume.skills or {},
                'additional': resume.additional or {}
            }
            ai_result = generate_professional_summary(resume_data)
            if ai_result.get('success'):
                return JsonResponse({'success': True, 'summary': ai_result['summary']})
            else:
                return JsonResponse({'success': False, 'error': ai_result.get('error', 'AI summary generation failed')}, status=500)
        except Exception as e:
            logger.error(f"AI summary generation error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to generate summary'}, status=500)
    return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)

def get_available_templates(request):
    """Get all available resume templates"""
    if request.method == 'GET':
        try:
            templates = templates_for_api()
            return JsonResponse({
                'success': True,
                'templates': templates,
                'count': len(templates)
            })
            
        except Exception as e:
            logger.error(f"Error getting templates: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to get templates'
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)

@login_required
def ai_resume_assistant(request):
    """AI Resume Assistant interface with chat and preview"""
    
    context = {
        'hero_content': {
            'page_title': 'AI Resume Assistant',
            'page_description': 'Create your perfect resume with our AI assistant. Chat naturally and see your resume update in real-time.'
        },
        'resume_templates': templates_for_gallery(),
        'resume_template_sections': get_resume_template_gallery_sections(),
        'default_template_id': DEFAULT_TEMPLATE_ID,
    }
    return render(request, 'resume_builder/resumeassistant.html', context)


def _prepare_message_with_smart_resume_context(message, user, _thread_id=None):
    """
    Determine when resume data is needed and include it only when necessary.
    """
    resume_keywords = [
        'resume', 'cv', 'curriculum vitae', 'create resume', 'build resume', 'edit resume',
        'update', 'change', 'modify', 'delete', 'remove', 'add', 'edit', 'save',
        'experience', 'education', 'skills', 'personal', 'summary', 'template',
        'university', 'college', 'school', 'degree', 'company', 'job', 'work',
        'position', 'title', 'employment', 'career', 'professional', 'work history',
        'academic', 'qualification', 'certification', 'training', 'course',
        'skill', 'competency', 'expertise', 'proficiency', 'knowledge',
        'project', 'achievement', 'accomplishment', 'responsibility', 'duty',
        'technology', 'software', 'tool', 'language', 'framework', 'platform'
    ]

    message_lower = message.lower()
    is_resume_related = any(keyword in message_lower for keyword in resume_keywords)

    print(f"🔍 Message analysis: '{message}'")
    print(f"🔍 Resume-related: {is_resume_related}")

    if not is_resume_related:
        print("📄 No resume data needed for general chat")
        return message

    try:
        from django.utils import timezone

        recent_time = timezone.now() - timedelta(minutes=10)
        current_resume = Resume.objects.filter(
            user=user,
            created_at__gte=recent_time
        ).order_by('-created_at').first()

        if current_resume:
            resume_data = {
                'resume_id': str(current_resume.id),
                'name': current_resume.name,
                'template_id': current_resume.template_id,
                'personal_info': current_resume.personal_info or {},
                'experience': current_resume.experience or [],
                'education': current_resume.education or [],
                'skills': current_resume.skills or {},
                'additional': current_resume.additional or {},
                'draft': current_resume.draft,
                'updated_at': current_resume.updated_at.isoformat() if current_resume.updated_at else None
            }

            enhanced_message = f"""
Current Resume State (as of {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}):
{json.dumps(resume_data, indent=2)}

User Request: {message}

Note: You can edit and delete entries using their array indices (0, 1, 2, etc.). The resume data above shows the current state.

CRITICAL: When editing or deleting entries, you MUST search through the current resume data provided above to find the correct entry. Do NOT rely on conversation history or previous entries - always use the current data shown above to determine the correct index.

SEARCH PROCESS FOR EDITING/DELETING:
1. First, identify what the user wants to edit/delete (institution name, company name, etc.)
2. Search through the current education/experience arrays in the resume data above
3. Look for EXACT or PARTIAL matches in the relevant fields
4. Use the ARRAY INDEX (0, 1, 2, etc.) of the matching entry
5. If multiple entries match, ask the user to be more specific
6. If no entry matches, inform the user that the specified entry was not found

EXAMPLE SEARCH PROCESS:
- User says: "update tech university to techno"
- Search education array for "tech university" in institution field
- Find entry at index 1 with institution "Tech University"
- Call edit_education with education_index=1, field="institution", value="Techno University"

CRITICAL: Before calling edit_education or delete_education, you MUST:
1. State which entry you found (e.g., "I found Tech University at index 1")
2. Confirm the action you're taking (e.g., "I will update Tech University to Techno University")
3. Then call the function with the correct index
"""
            print("📄 Resume data included for AI context")
            print(f"📄 Resume ID: {resume_data['resume_id']}")
            print(f"📄 Education entries: {len(resume_data['education'])}")
            print(f"📄 Experience entries: {len(resume_data['experience'])}")

            print("\n🔍 DEBUG: Current Education Entries:")
            for i, edu in enumerate(resume_data['education']):
                print(f"  Index {i}: {edu.get('degree', 'N/A')} from {edu.get('institution', 'N/A')}")
            print("🔍 DEBUG: End Education Entries\n")

            return enhanced_message

        print("📄 No recent resume found - treating as new conversation")
        return message
    except Exception as e:
        print(f"⚠️ Error getting resume data: {e}")
        return message


@login_required
@csrf_exempt
@require_http_methods(["POST"])
def chat_with_ai(request):
    """
    Django view for handling AI chat with proper user management
    
    This view integrates with the actual AI assistant manager to provide
    real resume building functionality through natural conversation.
    """
    print("\n" + "="*50)
    print("🔍 DEBUG: chat_with_ai view called")
    print("="*50)
    
    try:
        # Get the authenticated user
        user = request.user
        user_id = str(user.id)
        print(f"👤 User ID: {user_id}")
        print(f"👤 Username: {user.username}")
        
        # Parse the request
        print("📥 Parsing request body...")
        data = json.loads(request.body)
        message = data.get('message', '')
        thread_id = data.get('thread_id')
        print(f"💬 Message: {message}")
        print(f"🧵 Thread ID: {thread_id}")
        
        # Get or create assistant
        print("🤖 Getting or creating assistant...")
        manager, assistant_id = get_or_create_assistant()
        print(f"🤖 Assistant ID: {assistant_id}")
        
        if not assistant_id:
            print("⚠️ No assistant ID, returning fallback response")
            return JsonResponse({
                'success': True,
                'response': "I'm here to help you build your resume! I can help you create a new resume, edit existing ones, or get resume writing tips. What would you like to do?",
                'thread_id': thread_id or 'mock_thread_123',
                'user_id': user_id
            })
        
        # Create thread if not provided
        if not thread_id:
            print("🧵 Creating new thread...")
            thread_id = manager.create_thread()
            print(f"🧵 New thread ID: {thread_id}")
            if not thread_id:
                print("❌ Failed to create thread")
                return JsonResponse({
                    'success': False,
                    'error': 'Failed to create conversation thread'
                }, status=500)
        
        # Send message to AI with user context
        print("🚀 Sending message to AI...")
        print(f"📤 Query: {message}")
        print(f"👤 User ID: {user_id}")
        print(f"🧵 Thread ID: {thread_id}")
        print(f"🤖 Assistant ID: {assistant_id}")
        
        # Smart resume data inclusion - only include if we have a resume_id from previous AI interaction
        enhanced_message = _prepare_message_with_smart_resume_context(message, user, thread_id)
        
        result = manager.add_message_and_run(
            thread_id=thread_id,
            assistant_id=assistant_id,
            query=enhanced_message,
            user_id=user_id
        )
        
        print(f"📥 AI Result: {result}")
        
        if result:
            response_data = {
                'success': True,
                'response': result['response'],
                'thread_id': thread_id,
                'user_id': user_id
            }
            
            # Add resume ID if available
            if result.get('resume_id'):
                response_data['resume_id'] = result['resume_id']
                print(f"📄 Resume ID included: {result['resume_id']}")
            
            # Add action if available (from function results)
            if result.get('action'):
                response_data['action'] = result['action']
                print(f"🔧 Action included: {result['action']}")
            
            # Add action data if available
            if result.get('data'):
                response_data['data'] = result['data']
                print(f"📊 Action data included: {result['data']}")
            
            print("✅ Returning successful response")
            print(f"📤 Response data: {response_data}")
            return JsonResponse(response_data)
        else:
            print("❌ No result from AI")
            return JsonResponse({
                'success': False,
                'error': 'No response from assistant',
                'thread_id': thread_id
            }, status=500)
            
    except json.JSONDecodeError as e:
        print(f"❌ JSON decode error: {e}")
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        print(f"❌ Unexpected error: {str(e)}")
        import traceback
        print(f"📋 Traceback: {traceback.format_exc()}")
        return JsonResponse({
            'success': False,
            'error': 'Internal server error'
        }, status=500)

@login_required
def get_user_resumes(request):
    """
    Django view to get all resumes for the authenticated user
    
    This demonstrates how the AI assistant can access user-specific data
    """
    try:
        user = request.user
        user_id = str(user.id)
        
        # Get user's resumes
        resumes = user.resumes.all().order_by('-updated_at')
        
        resume_list = []
        for resume in resumes:
            resume_list.append({
                "resume_id": str(resume.id),
                "name": resume.name,
                "template_id": resume.template_id,
                "draft": resume.draft,
                "created_at": resume.created_at.isoformat() if resume.created_at else None,
                "updated_at": resume.updated_at.isoformat() if resume.updated_at else None
            })
        
        return JsonResponse({
            'success': True,
            'resumes': resume_list,
            'count': len(resume_list),
            'user_id': user_id
        })
        
    except Exception as e:
        logger.error(f"Error in get_user_resumes: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Internal server error'
        }, status=500)

@login_required
def get_resume_preview(request, resume_id):
    """
    Get resume preview HTML for the AI assistant interface
    
    This view returns the rendered HTML of a resume for the preview panel
    """
    try:
        # Get the resume
        resume = get_object_or_404(Resume, id=resume_id, user=request.user)
        
        # Get template from query params or use resume's default
        template_id = normalize_template_id(
            request.GET.get('template') or resume.template_id or None
        )
        
        # Prepare resume data for template rendering
        resume_data = {
            'personal_info': resume.personal_info or {},
            'experience': resume.experience or [],
            'education': resume.education or [],
            'skills': resume.skills or {},
            'additional': resume.additional or {}
        }
        
        # Render the resume with the chosen template
        html_content = _render_resume_template_html(
            request, resume_data, template_id, resume_id=resume.pk
        )
        
        return HttpResponse(html_content, content_type='text/html')
        
    except Exception as e:
        logger.error(f"Error in get_resume_preview: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Failed to load resume preview'
        }, status=500)


@login_required
def test_resumes(request):
    """Test resumes view"""
    resumes = request.user.resumes.all().order_by('-updated_at')
    return render(request, 'test/test_resumes.html', {'resumes': resumes})

@login_required
def resume_list_tab(request):
    """Serve the resume list tab content via AJAX"""
    resumes = Resume.objects.filter(user=request.user).order_by('-updated_at')
    context = {
        'resumes': resumes,
    }
    return render(request, 'resume_builder/component/resume_list_tab.html', context)

@login_required
def cover_letter_tab(request):
    """Serve the cover letter tab content via AJAX"""
    # TODO: Implement cover letter logic
    context = {
        'cover_letters': [],  # Placeholder
    }
    return render(request, 'resume_builder/component/cover_letter_tab.html', context)

@login_required
def templates_tab(request):
    """Serve the templates tab content via AJAX"""
    context = {
        'templates': templates_for_api(),
    }
    return render(request, 'resume_builder/component/templates_tab.html', context)

def resume_templates(request):
    """Display available resume templates with preview cards"""
    context = {
        'hero_content': {
            'page_title': 'Resume Templates',
            'page_description': 'Choose from our professional resume templates designed to help you stand out to employers.'
        },
        'resume_templates': templates_for_gallery(),
        'resume_template_sections': get_resume_template_gallery_sections(),
        'default_template_id': DEFAULT_TEMPLATE_ID,
    }
    return render(request, 'resume_templates/resume_templates.html', context)
